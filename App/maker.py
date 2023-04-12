import random
from docx import Document
import docx
import re
import csv
pattern_order = r'[0-9]'
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import json 

def combine_word_documents(files):
    merged_document = Document()
    
    for index, file in enumerate(files):
        sub_doc = Document(file)
        if index < len(files)-1:
           sub_doc.add_page_break()
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
    
    change_orientation(merged_document)
    merged_document.save('merged.docx')
    merged_document.save('merged.doc')
    merged_document.save('text.bin')

def change_orientation(document):
    for section in document.sections:
        section.orientation = docx.enum.section.WD_ORIENTATION.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
def addCheckbox(para, box_id, name,checked):
    run = para.add_run()
    tag = run._r
    fld = docx.oxml.shared.OxmlElement('w:fldChar')
    fld.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')

    ffData = docx.oxml.shared.OxmlElement('w:ffData')
    e = docx.oxml.shared.OxmlElement('w:name')
    e.set(docx.oxml.ns.qn('w:val'), 'Check1')
    ffData.append(e)
    ffData.append(docx.oxml.shared.OxmlElement('w:enabled'))
    e = docx.oxml.shared.OxmlElement('w:calcOnExit')
    e.set(docx.oxml.ns.qn('w:val'), checked)
    ffData.append(e)
    e = docx.oxml.shared.OxmlElement('w:checkBox')
    e.append(docx.oxml.shared.OxmlElement('w:sizeAuto'))
    ee = docx.oxml.shared.OxmlElement('w:default')
    ee.set(docx.oxml.ns.qn('w:val'), checked)
    e.append(ee)
    ffData.append(e)

    fld.append(ffData)
    tag.append(fld)

    run2 = para.add_run()
    tag2 = run2._r
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), str(box_id))
    start.set(docx.oxml.ns.qn('w:name'), name)
    tag2.append(start)

    run3 = para.add_run()
    tag3 = run3._r
    instr = docx.oxml.OxmlElement('w:instrText')
    instr.text = 'FORMCHECKBOX'
    tag3.append(instr)

    run4 = para.add_run()
    tag4 = run4._r
    fld2 = docx.oxml.shared.OxmlElement('w:fldChar')
    fld2.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
    tag4.append(fld2)

    run5 = para.add_run()
    tag5 = run5._r
    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), str(box_id))
    end.set(docx.oxml.ns.qn('w:name'), name)
    tag5.append(end)
def addParagraph(text,document):
    
    p  = document.add_paragraph(text.replace("\n"," "))

def addBullet(text,document,max):
    """styles = document.styles
    if('MyListNumberingStyle' in document.styles):
        num_format = document.styles['MyListNumberingStyle']
        print('test')
    else:
        num_format = document.styles.add_style('MyListNumberingStyle', WD_STYLE_TYPE.PARAGRAPH)  
    num_format.base_style = document.styles['List Number']
    num_format.paragraph_format.left_indent = Pt(36)
    num_format.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    num_fmt_elem = num_format._element.xpath('.//w:numFmt')
    if len(num_fmt_elem) > 0:
        num_fmt_elem[0].val = 'decimal'
    start_elem = num_format._element.xpath('.//w:start')
    if len(start_elem) > 0:
        start_elem[0].val = '1'"""
    i=1
    for part in text:
        p = document.add_paragraph(str(i)+'.'+' '+'\t'+part)
        i+=1
    
            
def addTable(document,text,rows,cols):
    ## yea idk why this is upside down i should probably fix that 
    table  = document.add_table(rows=cols,cols=rows)
    header = table.rows[0].cells
    for i in range (0,len(header)):
        p = header[i].paragraphs[0]
        p.add_run(text[0][i]).bold=True
    for i in range(1,cols):
        for j in range(0,rows):
            table.cell(i,j).text = text[i][j]
    table.style='Table Grid'
        
'''
Open the json
take out parts 
'''
def writeDoc(data):
    document = Document()
    
    data = json.load(f)
    part1 = data['page1']
    part2 = data['page2']
    part3 = data['page3']
    part4 = data['page4']
    part5 = data['page5']
    part6 = data['page6']
    part7 = data['page7']
    part8 = data['page8']
    part9 = data['page9']
    bigPart = {**part1,**part2,**part3}
    i=0
    p = document.add_paragraph()
    p.add_run("THE UNIVERSITY OF THE WEST INDIES DRAFT \n PROPOSAL FOR NEW/REVISED UNDERGRADUATE COURSE").bold=True
    with open('CourseText.txt') as courseInfo:
        data1 = []
        line = courseInfo.readlines()
        for key in bigPart:
            value = bigPart[key]
            data1.append(value)
        for part in line:
            if 'Credits' in part:
                ## the next thing is a table lol
                part = part.strip('\n')
                p  = document.add_paragraph()
                p.add_run(part).bold = True
                p.add_run(data1[i]) 
                header = ['Projected Enrolment','Year 1','Year 2','Year 3','Other Year(s)']
                tabledata = []
                tabledata.append(header)
                rowdata = []
                partList = list(part2.items())
                j=0
                key, value = partList[j]
                rowdata.append('Full Time Student')
                while 'Full Time' in key:
                    rowdata.extend(value)
                    j=j+1
                    key, value = partList[j]
                tabledata.append(rowdata)
                rowdata = rowdata.copy()
                rowdata.clear()
                rowdata.append('Part Time Student')
                while 'Part Time' in key and j<len(partList):
                    key, value = partList[j]
                    rowdata.extend(value)
                    j=j+1
                tabledata.append(rowdata)
                addTable(document=document,text=tabledata,rows=len(tabledata[0]),cols=len(tabledata))
                i=i+len(part2)
            part = part.strip('\n')
            p  = document.add_paragraph()
            p.add_run(part).bold = True
            p.add_run(data1[i])
            i+=1 
        p  = document.add_paragraph()
        p.add_run('Course Description').bold = True
        addParagraph(data1[i],document=document)
        i=i+1
        p  = document.add_paragraph()
        p.add_run('Course Rationale').bold = True
        addParagraph(data1[i],document=document)
        i=i+1

        p  = document.add_paragraph()
        p.add_run('Course Aims').bold = True
        value = data1[i]
        if str(1)+"." in value:
            ##if a 1 is here that means that there is most likely a bullter point the user was trying to do a bullet point 
            newList = data1[i].split("\n")
            p  = document.add_paragraph()
            p.add_run(newList[0])
            for i in range(1,len(newList)):
                newList[i] = newList[i].replace(str(i)+".","")
            addBullet(newList,document=document,max=len(newList))
        else:
            addParagraph(data1[i],document=document)
        partList = list(part4.items())
        line = []
        i=0
        key, value = partList[i]
        while 'courseContent' not in key:
            if value != "None":
                line.append(value)
            i=i+1
            key, value = partList[i]
        p = document.add_paragraph()
        p.add_run('Learning Outcomes').bold = True
        p = document.add_paragraph()
        p.add_run('Upon the successful completion of this course, the student will be able to:')
        addBullet(line,document=document,max=len(line))
        ## so when we putting this into the system we should probably have a check based on the person cause it kinda look like it have to stay the same
        ##but for now i am leaving it as a point
        line = partList[i:]
        p = document.add_paragraph()
        p.add_run('Course Content').bold = True
       
        key, value = partList[i]
        points=[]
        while 'teachingMethods' not in key:
            points.append(value)
            i=i+1
            key, value = partList[i]
        if(len(points)>1):
            p = document.add_paragraph()
            p.add_run('The following main topics are covered in this course:')
            addBullet(points,document=document,max=len(points))
        else:
            addParagraph(points[0],document=document)

        key, line = partList[i]
        p = document.add_paragraph()
        p.add_run('Teaching methods').bold = True
        addParagraph(line,document=document)

        p = document.add_paragraph()
        p.add_run('Contact and credits hours').bold = True
        header = ['Type','Duration','Contact Hours','Credit Hours']
        tabledata = []
        tabledata.append(header)
        rowdata = []
        partList = list(part5.items())
        i=0

        key, value = partList[i]
        rowdata.append('Lecture')
        while 'lecture' in key:
            rowdata.append(value)
            i=i+1
            key, value = partList[i]
            
        tabledata.append(rowdata)
        rowdata = rowdata.copy()
        rowdata.clear()
        rowdata.append('Tutorial')
        while 'tutorial' in key :
            rowdata.append(value)
            i=i+1
            key,value = partList[i]
            
        tabledata.append(rowdata)
        rowdata = rowdata.copy()
        rowdata.clear()
        rowdata.append('lab')

        while 'lab' in key :
            rowdata.append(value)
            i=i+1
            key,value = partList[i]
            
        tabledata.append(rowdata)
        rowdata = rowdata.copy()
        rowdata.clear()
        rowdata.append('other')
        while 'other' in key:
            rowdata.append(value)
            i=i+1
            key,value = partList[i]
        tabledata.append(rowdata)
        rowdata = rowdata.copy()
        rowdata.clear()
        rowdata.append('Total')
        while 'total' in key and i<len(partList):
            key, value = partList[i]
            rowdata.append(value)
            i=i+1 
        tabledata.append(rowdata)
        addTable(document=document,text=tabledata,rows=len(tabledata[0]),cols=len(tabledata))
        i=0
        dataList = list(part6.items())
        key, value = dataList[i]
        p = document.add_paragraph()
        p.add_run('Assessment Description').bold = True
        addParagraph(value,document=document)

        p = document.add_paragraph()
        p.add_run('Course Assessment Type and Course Learning Outcome Matrix').bold = True
        assignment = ['Assessment', 'Learning Outcomes' ,'Weighting% ','Assessment Description' ,'Duration']
        i=0
        learningOutcomesNum = len(part4)-2
        heading_text = assignment
    
        for i in range(1,learningOutcomesNum+1):
            heading_text.append(str(i))
        rowlen = int((len(part6)-1)/5)
        table1 = document.add_table(rows=rowlen+2,cols=len(heading_text)-1,style='Table Grid')
        header = table1.rows[0].cells
        numbers = table1.rows[1].cells
        j=0
        i=0
        outcomesNum = len(header) - 4
        while(j<len(header)):
            if(heading_text[i]=='Learning Outcomes'):
                header[j].text = heading_text[i]
                j=j+learningOutcomesNum-1
            elif(heading_text[i].isnumeric()):
                break
            else:
                header[j].text = heading_text[i]
            j=j+1
            i=i+1
        for i in range(1,outcomesNum):
            table1.cell(0,i).merge(table1.cell(0,i+1))
        for i in range(1,outcomesNum+1):
            table1.cell(1,i).text = heading_text[i+4]
        for i in range(0,len(numbers)):
            if table1.cell(1,i).text == '':
                table1.cell(1,i).merge(table1.cell(0,i))
        partList = list(part6.items())
        partList = list(part6.items())
        k=1
        key,value = partList[k]
        last = 0
        j=0
        for i in range(0,int((len(partList)-1)/5)):
            j=0
            while j<len(assignment)-1:
                if 'LearningOutcomes' in key:
                    for part in value:
                        if(part.isnumeric()):
                            table1.cell(i+2,j+int(part)-1).text = 'X'
                            last = int(part)
                    while table1.cell(0,j).text != 'Weighting% ':
                        j=j+1
                    j=j-1
                else:
                    table1.cell(i+2,j).text = value
                if k<len(partList)-1:
                    k=k+1
                key ,value= partList[k]
                j=j+1
        p = document.add_paragraph()
        p.add_run('Course Calendar').bold = True
        partList = list(part7.items())
        i=0
        header = ['Week', 'Topic', 'Required Reading Learning Resources', 'Learning Activities', 'Assignments', '']
        header1= ['', '', '', '', 'Name', 'Due Date']
        tabledata = []
        tabledata.append(header)
        tabledata.append(header1)
        rowdata = []
        key, value = partList[i]
        match = re.search(r'\d{1,2}', key)
        rowdata.append(str(match.group()))
        while i<len(partList)-1:
            key, value = partList[i]
            match = re.search(r'\d{1,2}', key)
            rowdata.append(value)
            i=i+1
            key, value = partList[i]
            match1 = re.search(r'\d{1,2}', key)
            if match.group() != match1.group(): ## it's a new week
                tabledata.append(rowdata)
                rowdata = rowdata.copy()
                rowdata.clear()
                rowdata.append(str(match1.group()))
        rowdata.append(value)      
        tabledata.append(rowdata)
        rowdata = rowdata.copy()
        rowdata.clear()
        addTable(document,tabledata,len(tabledata[0]),len(tabledata))
        partList = list(part8.items())
        p = document.add_paragraph()
        p.add_run('Attributes of the Ideal UWI Graduate').bold = True
        for key, value in partList:
            p = document.add_paragraph()
            addCheckbox(p,random.randint(16*1024, 32*1024),'zane',str(int(value)))
            p.add_run(key)
        partList = list(part9.items())
        data2 =''
        i=0
        key, value = partList[i]
        while key.find('Resources') !=-1:
            line = key+'-'+value +' \n'
            data2 = data2+line
            i=i+1
            key, value = partList[i]
        p = document.add_paragraph()
        p.add_run('Resources').bold = True
        addParagraph(data2,document=document)
        key, value = partList[i]
        p = document.add_paragraph()
        p.add_run('Staffing Requirements').bold = True
        addParagraph(value,document=document)
    courseInfo.close()

    ################################
    change_orientation(document)
    document1 = Document('TemplateFile.docx')
    change_orientation(document1)
    files = ['even more testing.docx','TemplateFile.docx']
    document.save("even more testing.docx")
    combine_word_documents(files)
